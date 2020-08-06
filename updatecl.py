from docx import Document
from datetime import date

today = date.today()

document = Document('cover_letter_template.docx')

company_name =  str(raw_input("Enter the name of the company you are applying for:\n"))
company_adress = str(raw_input("Enter the adress of the company you are applying for:\n"))
position_name = str(raw_input("Enter the name of the position you are applying for:\n"))
position_posting_location = str(raw_input("Where did you find the job posting?:\n"))
print 'Available Bodies:\n'
for paragraph in document.paragraphs:
    if '<' in paragraph.text and '</' not in paragraph.text:
        print paragraph.text+'\n'

body_number = str(raw_input("Which preset cover letter body would you like to use?:(enter only the number ei. 1,2,3...)\n"))
delete_mode = False
keep_mode = False

for paragraph in document.paragraphs:
    if '[COMPANY NAME]' in paragraph.text:
        paragraph.text = {company_name}
    if '[COMPANY ADRESS]' in paragraph.text:
    	paragraph.text = {company_adress}
    if '[POSITION NAME]' in paragraph.text:
    	paragraph.text = paragraph.text.replace("[POSITION NAME]",position_name)
    if '[POSITION POSTING LOCATION]' in paragraph.text:
    	paragraph.text = paragraph.text.replace("[POSITION POSTING LOCATION]",position_posting_location)
    if '[TODAYS DATE]' in paragraph.text:
        paragraph.text = {today.strftime("%B %d, %Y")}
    if '<body' in paragraph.text:
    	if '<body '+body_number in paragraph.text:
    		keep_mode = True
    		paragraph._element.getparent().remove(paragraph._element)
    	else:
	    	paragraph._element.getparent().remove(paragraph._element)
	    	delete_mode = True
    elif '</body' in paragraph.text:
    	if '<body '+body_number in paragraph.text:
    		keep_mode = False
    		paragraph._element.getparent().remove(paragraph._element)
    	else:
	    	paragraph._element.getparent().remove(paragraph._element)
	    	delete_mode = False
    elif delete_mode:
    		paragraph._element.getparent().remove(paragraph._element)
    
document.save('cover_letter.docx')

